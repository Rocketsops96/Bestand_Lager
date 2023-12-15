from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL



data_table1 = None
def insert_data_into_tables(input_file, output_file, data_table1):
    doc = Document(input_file)
    try:
        table1 = doc.tables[0]  # Первая таблица

    except IndexError:
        print("Ошибка: Не удалось найти одну из таблиц.")
        return

    # Вставляем данные в первую таблицу

    for row_num, row_data in enumerate(data_table1):
        for col_num, cell_data in enumerate(row_data):
            cell = table1.cell(row_num, col_num)
            if not cell.text.strip():
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.clear()

                cell.text = str(cell_data)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Сохраняем изменения
    doc.save(output_file)
