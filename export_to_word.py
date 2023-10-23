import sqlite3
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import main

# Подключитесь к базе данных BAU
conn = sqlite3.connect("bau.db")
cursor = conn.cursor()

# Выполните SQL-запрос для получения данных из базы данных
cursor.execute("SELECT * FROM bvhgf")  # Замените "ваша_таблица" на имя вашей таблицы
data = cursor.fetchall()

# Создайте новый документ Word
doc = Document()

# Добавьте заголовок
title = doc.add_heading("Materialliste", level=1)
title2 = doc.add_heading("Verkehrssicherung", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Создайте таблицу для данных
table = doc.add_table(rows=1, cols=4)  # Задайте количество строк и столбцов
table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
table.autofit = False  # Отключаем автонастройку ширины столбцов

# Установите ширину столбцов
for col in table.columns:
    col.width = Pt(100)  # Здесь можно настроить ширину столбцов по вашему усмотрению

# Добавьте заголовки столбцов
table.rows[0].cells[0].text = "Заголовок 1"
table.rows[0].cells[1].text = "Заголовок 2"
table.rows[0].cells[2].text = "Заголовок 3"
table.rows[0].cells[3].text = "Заголовок 4"

# Заполните таблицу данными из базы данных
for row_data in data:
    row = table.add_row().cells
    row[0].text = str(row_data[0])  # Замените на соответствующий индекс поля
    row[1].text = str(row_data[1])  # Замените на соответствующий индекс поля
    row[2].text = str(row_data[2])  # Замените на соответствующий индекс поля
    row[3].text = str(row_data[3])  # Замените на соответствующий индекс поля

# Сохраните документ Word
doc.save("отчет_bau.docx")

# Закройте соединение с базой данных
conn.close()
