from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import regbase 
from datetime import datetime

data_table2 = None
data_table1 = None
def insert_data_into_tables(input_file, output_file, data_table1, data_table2):
    # Открываем существующий документ Word
    doc = Document(input_file)

    # Выбираем нужные таблицы (предположим, что они находятся в документе)
    try:
        table1 = doc.tables[0]  # Первая таблица
        table2 = doc.tables[1]  # Вторая таблица
    except IndexError:
        print("Ошибка: Не удалось найти одну из таблиц.")
        return

    # Вставляем данные в первую таблицу
    for row_num, row_data in enumerate(data_table1):
        for col_num, cell_data in enumerate(row_data):
            cell = table1.cell(row_num, col_num)
            if not cell.text.strip():
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.clear()
                cell.text = str(cell_data)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Вставляем данные во вторую таблицу
    for row_num, row_data in enumerate(data_table2):
        for col_num, cell_data in enumerate(row_data):
            cell = table2.cell(row_num, col_num)
            if not cell.text.strip():
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.clear()
                cell.text = str(cell_data)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Сохраняем изменения
    doc.save(output_file)

if __name__ == "__main__":
    input_document = "Stundenbericht Verkehrssicherung.docx"
    output_document = "test.docx"
    
    conn = regbase.create_conn()
    cursor = conn.cursor()
    cursor.execute("SELECT id, name_bau, kostenstelle_vvo, bauvorhaben, ort, strasse, ausfurung_von, ausfurung_bis, vrao_ab, vrao_bis, ansprechpartner, status FROM bau WHERE id = 7")
    data = cursor.fetchone()
    ccussor = conn.cursor()
    bau = data[2]
    ccussor.execute("SELECT id, name_capo, bau, action, date_ab, date_bis FROM sicherung WHERE bau = %s ",(bau,))
    sicherung = ccussor.fetchone()
    time_start_str = sicherung[4]
    time_end_str = sicherung[5]
    
    date = sicherung[4]
    datetime_with_time = datetime.strptime(date, '%Y-%m-%d %H:%M:%S')
    date_withot_time = datetime_with_time.date()
    formated_date = date_withot_time.strftime('%d.%m.%Y')
    # Преобразование строк в объекты datetime
    time_start = datetime.strptime(time_start_str, '%Y-%m-%d %H:%M:%S')
    time_end = datetime.strptime(time_end_str, '%Y-%m-%d %H:%M:%S')
    print(time_start)
    print(time_end)
    # Вычисление разницы во времени
    time_difference = time_end - time_start

    # Извлечение разницы в часах
    hours_difference = time_difference.total_seconds() / 3600
    rounded_difference = round(hours_difference * 2) / 2
    print(f"Разница в часах: {rounded_difference}")
    data_table1 = [
        ["", data[2], "", "", f"{data[6]} - {data[7]}"],
        ["", "", "", "", ""],
        ["", data[3], "","",data[8]],
        ["", f"{data[4]}, {data[5]}","","", data[9]],
    ]

    # Пример данных для второй таблицы
    data_table2 = [
        ["", "", "", "", "", ""],
        ["", "", "", "", "", "","", "", "", "", "", ""],
        [1, f"{sicherung[1]}", f"{formated_date}", "", "", "","", "", "", "", "", f"{rounded_difference}"],
        [2, "", "", "", "", "","", "", "", "", "", ""],
        [3, "", "", "", "", "","", "", "", "", "", ""],
        [4, "", "", "", "", "","", "", "", "", "", ""],
        [5, "", "", "", "", "","", "", "", "", "", ""],
        [6, "", "", "", "", "","", "", "", "", "", ""],
        [7, "", "", "", "", "","", "", "", "", "", ""],
        [8, "", "", "", "", "","", "", "", "", "", ""],
        [9, "", "", "", "", "","", "", "", "", "", ""],
        [10, "", "", "", "", "","", "", "", "", "", ""],
        [11, "", "", "", "", "","", "", "", "", "", ""],
        [12, "", "", "", "", "","", "", "", "", "", ""],
        [13, "", "", "", "", "","", "", "", "", "", ""],
        [14, "", "", "", "", "","", "", "", "", "", ""],
        [15, "", "", "", "", "","", "", "", "", "", ""],
        [16, "", "", "", "", "","", "", "", "", "", ""],
        ["", "", "", "", "", "","", "", "", "", "", ""],
    ]

    insert_data_into_tables(input_document, output_document, data_table1, data_table2)
